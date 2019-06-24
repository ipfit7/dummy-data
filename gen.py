from docx import Document
from factories.identity_factory import IdentityFactory
from models.identity_model import IdentityModel
from factories.email_factory import EmailFactory
from models.email_model import EmailModel
from repositories.treatment_repository import TreatmentRepository
import random
import mailbox


def create_doc(identity: IdentityModel) -> None:
    # Placeholder
    patient_nr = random.randint(10000, 99999)

    #create document
    document = Document()
    document.add_heading('Patient Dossier van ' + identity.full_name, 0)

    # Create table
    patientInfo = document.add_table(rows=5, cols=2)

    desc_cells = patientInfo.columns[0].cells
    desc_cells[0].text = "Patient nummer"
    desc_cells[1].text = "Naam"
    desc_cells[2].text = "Geboorte datum"
    desc_cells[3].text = "Woonplaats"
    desc_cells[4].text = "BSN"

    info_cells = patientInfo.columns[1].cells

    info_cells[0].text = str(patient_nr)
    info_cells[1].text = identity.full_name
    info_cells[2].text = str(identity.dateOfBirth)
    info_cells[3].text = identity.placeOfResidence
    info_cells[4].text = str(identity.bsn)

    document.save("dossier {0}.docx".format(patient_nr))

def create_mailbox(email: EmailModel) -> None:
    mbox = mailbox.mbox("mailbox.mbox")
    mbox.lock()
    try:
        msg = mailbox.mboxMessage()
        msg.set_unixfrom('author Sat Feb  7 01:05:34 2009')
        msg['From'] = email.email_from
        msg['To'] = email.email_to
        msg['Subject'] = email.subject
        msg.set_payload(email.content)
        mbox.add(msg)
        mbox.flush()
    finally:
        mbox.unlock()

if __name__ == "__main__":
    identity_factory = IdentityFactory()
    email_factory = EmailFactory()
    treatment_repos = TreatmentRepository()
    identities = [identity_factory.get_random_identity() for x in range(0, 1000)]

    for i in range(1, 20):
        patient = random.choice(identities)
        doctor = random.choice(identities)
        # create_doc(patient)
        create_mailbox(email_factory.create_email_from_treatment(doctor, patient, treatment_repos.get_random_treatment()))
