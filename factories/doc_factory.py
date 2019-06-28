from util.singleton import Singleton
from models.identity_model import IdentityModel
from factories.identity_factory import IdentityFactory
from repositories.photo_repository import PhotoRepository
import random
from docx import Document

class DocFactory(metaclass=Singleton):
    def create_doc(self, identity: IdentityModel, treatments: []) -> None:
        # Placeholder
        patient_nr = random.randint(10000, 99999)

        #create document
        document = Document()
        document.add_heading('Patient Dossier van ' + identity.full_name, 0)

        # Create table
        patientInfo = document.add_table(rows=5, cols=2)

        desc_cells = patientInfo.columns[0].cells
        desc_cells[0].text = "Patient nummer"
        desc_cells[1].text = "Naam"
        desc_cells[2].text = "Geboorte datum"
        desc_cells[3].text = "Woonplaats"
        desc_cells[4].text = "BSN"

        info_cells = patientInfo.columns[1].cells

        info_cells[0].text = str(patient_nr)
        info_cells[1].text = identity.full_name
        info_cells[2].text = str(identity.dateOfBirth)
        info_cells[3].text = identity.placeOfResidence
        info_cells[4].text = str(identity.bsn)

        document.add_heading('Behandelingen', 1)

        behandelingen = document.add_table(rows=len(treatments), cols=3)
        xray = False
        i = 0 
        for treatment in treatments:
            behandelingen.columns[0].cells[i].text = treatment.treatmentDesc
            behandelingen.columns[1].cells[i].text = "€{:.2f}".format(treatment.treatmentCost)
            behandelingen.columns[2].cells[i].text = str(treatment.treatmentDate)
            if treatment.treatmentDesc == "Xray":
                xray = True
            i += 1
        
        if xray:
            document.add_heading("Röntgenfoto's", 1)
            document.add_picture(PhotoRepository().get_random_photo())

        return document
